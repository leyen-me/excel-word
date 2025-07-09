import os
import re
import sys
import logging
from openpyxl import load_workbook
from docx import Document
import pyexcel

try:
    import colorama
    colorama.init()
    COLORAMA_LOADED = True
except ImportError:
    COLORAMA_LOADED = False

# 日志格式和颜色
class LogFormatter(logging.Formatter):
    COLORS = {
        'DEBUG': '\033[36m',    # 青色
        'INFO': '\033[32m',     # 绿色
        'WARNING': '\033[33m',  # 黄色
        'ERROR': '\033[31m',    # 红色
        'CRITICAL': '\033[41m', # 红底
    }
    RESET = '\033[0m'

    def format(self, record):
        msg = super().format(record)
        if COLORAMA_LOADED and record.levelname in self.COLORS:
            msg = f"{self.COLORS[record.levelname]}{msg}{self.RESET}"
        return msg

logger = logging.getLogger("excel-word")
logger.setLevel(logging.DEBUG)
ch = logging.StreamHandler()
formatter = LogFormatter('[%(levelname)s] %(message)s')
ch.setFormatter(formatter)
logger.handlers.clear()
logger.addHandler(ch)

if getattr(sys, 'frozen', False):
    BASE_DIR = os.path.dirname(sys.executable)
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

excel_folder = os.path.join(BASE_DIR, 'assets')
word_template_path = os.path.join(BASE_DIR, 'template', 'temp.docx')
output_folder = os.path.join(BASE_DIR, 'output')

def gen_cells_row(start_col, end_col, row):
    def col_range(start, end):
        for c in range(ord(start), ord(end) + 1):
            yield chr(c)
    return [f"{col}{row}" for col in col_range(start_col, end_col)]

def gen_cells(col, start_row, end_row):
    return [f"{col}{row}" for row in range(start_row, end_row + 1)]

fields = {
    '工程部位': 'R10',
    '设计强度等级': 'J15',
    '墩柱': 'A21',
    '测区平均值': gen_cells('S', 21, 30),
    '测区声速代表值': gen_cells('Z', 21, 30),
    '平测声速':'Z16',
    '修正为对测声速': gen_cells('AA', 21, 30),
    '测区强度代表值': gen_cells('AB', 21, 30),
    '构件强度推定值': 'AD21',
    '设计抗压强度等级': 'J15',
    '标准差': 'AC21',
    '平均值': gen_cells('AB', 21, 30)
}

decimal_places_map = {
    '测区平均值': 1,
    '测区声速代表值': 2,
    '修正为对测声速': 2,
    '测区强度代表值': 1,
    '平测声速': 3,
    '构件强度推定值': 1,
    '标准差': 3,
    '平均值': 1
}

def extract_cell_values(sheet, cell_ref):
    if isinstance(cell_ref, str) and ':' in cell_ref:
        # 区间
        cells = sheet[cell_ref]
        # cells 可能是二维的，需展开
        return [cell.value for row in cells for cell in row]
    elif isinstance(cell_ref, list):
        return [sheet[c].value for c in cell_ref]
    else:
        return sheet[cell_ref].value

def format_value_by_rule(key, value):
    rule = decimal_places_map.get(key)
    if key == '平测声速':
        if value == '/' or value == '／':
            return 1
        if rule and value not in (None, '', ' '):
            return f"{float(value):.{rule}f}"
        return value
    if isinstance(value, list):
        return [
            f"{float(v):.{rule}f}" if rule and v not in (None, '', ' ') else v
            for i, v in enumerate(value)
        ]
    else:
        if rule and value not in (None, '', ' '):
            return f"{float(value):.{rule}f}"
        return value

def trans_xls_to_xlsx(folder):
    for file in os.listdir(folder):
        if file.endswith('.xls'):
            xls_path = os.path.join(folder, file)
            xlsx_path = xls_path + 'x'
            logger.info(f"正在转换: {xls_path} -> {xlsx_path}")
            pyexcel.save_book_as(file_name=xls_path, dest_file_name=xlsx_path)
            logger.info(f"已转换: {file} -> {os.path.basename(xlsx_path)}")

def extract_data_from_excel(excel_path):
    logger.info(f"正在读取Excel: {excel_path}")
    wb = load_workbook(excel_path, data_only=True)
    data = {}
    for sheet in wb.worksheets:
        logger.info(f"处理Sheet: {sheet.title}")
        sheet_data = {}
        for key, cell in fields.items():
            value = extract_cell_values(sheet, cell)
            # 对“平均值”特殊处理
            if key == '平均值':
                # 过滤掉空值和非数字
                nums = [float(v) for v in value if v not in (None, '', ' ') and isinstance(v, (int, float, str)) and str(v).replace('.', '', 1).isdigit()]
                avg = sum(nums) / len(nums) if nums else ''
                value = format_value_by_rule(key, avg)
            else:
                value = format_value_by_rule(key, value)
            sheet_data[key] = value
            logger.debug(f"提取字段: {key}，单元格: {cell}，值: {value}")
            if value == '' or value == ' ':
                logger.warning(f"数据可能出错 {sheet.title} {key}")
                raise Exception(f"数据可能出错 {sheet.title} {key}")
        data[sheet.title] = sheet_data
    logger.info(f"完成Excel数据提取: {excel_path}")
    return data

def replace_placeholder_in_paragraph(paragraph, data):
    full_text = ''.join(run.text for run in paragraph.runs)
    # for key, value in data.items():
    #     placeholder = f"{{{{{key}}}}}"
    #     if placeholder in full_text:
    #         logger.debug(f"替换占位符: {placeholder} -> {value}")
    #     full_text = full_text.replace(placeholder, str(value))
    # if paragraph.runs:
    #     paragraph.runs[0].text = full_text
    #     for run in paragraph.runs[1:]:
    #         run.text = ''
    # 先处理带数字的占位符，如 {{测区平均值1}}
    def replace_match(match):
        key = match.group(1)
        idx = match.group(2)
        if key in data and isinstance(data[key], list):
            try:
                # idx是1起始，Python下标是0起始
                _data = data[key][int(idx) - 1]
                if _data is None:
                    _data = "/"
                return str(_data)
            except (IndexError, ValueError):
                return ''
        # 如果不是数组，或者key不存在，返回原样
        return match.group(0)

    # 替换所有 {{字段名数字}} 占位符
    full_text = re.sub(r'\{\{([\u4e00-\u9fa5A-Za-z_]+?)(\d+)\}\}', replace_match, full_text)

    # 再处理普通的 {{字段名}} 占位符
    for key, value in data.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in full_text:
            logger.debug(f"替换占位符: {placeholder} -> {value}")
        if isinstance(value, list):
            value_str = ', '.join(str(v) for v in value)
        else:
            _data = value
            if _data is None:
                _data = "/"
            value_str = str(_data)
        full_text = full_text.replace(placeholder, value_str)

    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ''
def fill_word_template(data, output_path):
    logger.info(f"正在填充Word模板: {output_path}")
    doc = Document(word_template_path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholder_in_paragraph(paragraph, data)
    for paragraph in doc.paragraphs:
        replace_placeholder_in_paragraph(paragraph, data)
    doc.save(output_path)
    logger.info(f"Word已保存: {output_path}")

def clean_filename(name):
    return re.sub(r'[\\/:#*?"<>|]', '_', name)

def main():
    logger.info("程序启动")
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        logger.info(f"创建输出文件夹: {output_folder}")
    else:
        logger.info(f"输出文件夹已存在: {output_folder}")
        
    logger.info(f"开始转换xls为xlsx，目录: {excel_folder}")
    trans_xls_to_xlsx(excel_folder)
    
    for filename in os.listdir(excel_folder):
        if filename.endswith('.xlsx'):
            _filename = filename.replace(".xlsx", "")
            excel_path = os.path.join(excel_folder, filename)
            logger.info(f"处理Excel文件: {excel_path}")
            try:
                datas = extract_data_from_excel(excel_path)
            except Exception as e:
                logger.error(f"读取Excel数据失败: {excel_path}，错误信息: {e}")
                continue
            
            for sheet_name in datas:
                safe_sheet_name = clean_filename(sheet_name)
                word_name = f"{_filename}_{safe_sheet_name}.docx"
                output_path = os.path.join(output_folder, word_name)
                try:
                    fill_word_template(datas[sheet_name], output_path)
                    logger.info(f"已生成: {output_path}")
                except Exception as e:
                    logger.error(f"生成Word失败: {output_path}，错误信息: {e}")

#  python3 -m PyInstaller -F main.py --hidden-import=pyexcel_io.writers --hidden-import=pyexcel_xls --hidden-import=pyexcel_xlsx

if __name__ == '__main__':
    try:
        main()
        logger.info("程序执行完毕，按回车键退出。")
        input()
    except Exception as e:
        logger.critical(f"发生错误：{e}")
        import traceback
        traceback.print_exc()
        input("按回车键退出。")